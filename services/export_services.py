"""
Export services
===============
Converts the rich-text HTML produced by ``DocumentTab.get_content_html()``
(a QTextEdit's ``toHtml()`` output, with images already embedded as base64
data URIs) into two portable formats:

- ``html_to_markdown``   -> a plain ``.md`` string
- ``save_html_as_docx``  -> a real Word ``.docx`` file on disk

Both entry points also happily accept plain hand-written semantic HTML
(``<h1>``, ``<strong>``, ``<a href="">`` etc.), which is what the unit tests
exercise directly - that keeps the converters honest and testable without
needing a running Qt application.

Design
------
Qt's ``toHtml()`` does not use semantic tags for the things this app treats
as "styles": a heading is just a ``<p>`` with extra top/bottom margin whose
text is wrapped in a bold, oversized ``<span>``, and a code block is a
``<span>`` in a monospace font.  ``_normalize_qt_html`` rewrites that markup
into plain semantic HTML (``<h1>``, ``<code>``, ``<strong>`` ...) first, so
the rest of the pipeline only has to deal with ordinary tags - whether they
came from Qt or from a caller's own HTML string.
"""

from __future__ import annotations

import base64
import re
from html import unescape
from html.parser import HTMLParser
from io import BytesIO
from pathlib import Path
from typing import Optional

# ============================================================================
# Shared HTML normalization (Qt's span-based "styles" -> semantic tags)
# ============================================================================

# font-size (pt) -> heading level, mirrors FormattingController.apply_text_style
_HEADING_SIZES = {28: 1, 24: 1, 20: 2, 18: 3, 16: 4, 14: 5, 12: 6}
_MONOSPACE_FAMILIES = ("courier", "consolas", "monospace", "menlo", "monaco")

# Common web-safe / "basic" fonts. A font-family is only ever carried through
# to the export (as a Markdown <font> tag or a real Word font) when it's in
# this set; anything else falls back to the document's normal font, since we
# can't know if the reader's machine even has it installed.
_BASIC_HTML_FONTS = {
    "arial", "helvetica", "verdana", "tahoma", "trebuchet ms",
    "times new roman", "times", "georgia", "garamond", "book antiqua",
    "palatino linotype", "palatino", "comic sans ms", "impact",
    "arial black", "calibri", "cambria", "lucida console",
    "lucida sans unicode", "segoe ui",
}

# Qt renders hyperlinks as blue + underlined text; this lets us recognize
# that combination and avoid wrapping link text in a redundant color span.
_LINK_BLUE = "#0000ff"

_NAMED_COLORS = {
    "black": "000000", "white": "ffffff", "red": "ff0000", "green": "008000",
    "blue": "0000ff", "yellow": "ffff00", "orange": "ffa500", "purple": "800080",
    "pink": "ffc0cb", "gray": "808080", "grey": "808080", "brown": "a52a2a",
    "cyan": "00ffff", "magenta": "ff00ff",
}

_STYLE_ATTR_RE = re.compile(r'<span\s+style="([^"]*)">(.*?)</span>', re.IGNORECASE | re.DOTALL)
_HEADING_BLOCK_RE = re.compile(
    r'<p\s+style="[^"]*margin-top:\s*(\d+)px;\s*margin-bottom:\s*(\d+)px[^"]*"[^>]*>'
    r'\s*<span\s+style="([^"]*)">(.*?)</span>\s*</p>',
    re.IGNORECASE | re.DOTALL,
)


def _style_dict(style: str) -> dict:
    """Parse a ``key:value; key:value`` inline CSS string into a dict."""
    result = {}
    for chunk in style.split(";"):
        if ":" not in chunk:
            continue
        key, _, value = chunk.partition(":")
        result[key.strip().lower()] = value.strip().lower()
    return result


def _heading_level_for_size(size_pt: int) -> Optional[int]:
    """Map a font size to the closest heading level this app produces."""
    if size_pt in _HEADING_SIZES:
        return _HEADING_SIZES[size_pt]
    if size_pt > 24:
        return 1
    # Snap to the nearest known size below it (e.g. 22pt -> H2's 20pt bucket)
    candidates = sorted(_HEADING_SIZES, reverse=True)
    for known in candidates:
        if size_pt >= known:
            return _HEADING_SIZES[known]
    return None


def _is_basic_font(family: str) -> bool:
    """Whether *family* is a common, web/Word-safe font we're willing to keep."""
    return family.strip().strip("'\"").lower() in _BASIC_HTML_FONTS


def _to_hex(value: str) -> Optional[str]:
    """Normalize a CSS color (hex, 3-digit hex, name, or rgb()) to 'RRGGBB'."""
    value = value.strip().lower()
    if value.startswith("#"):
        hexpart = value[1:]
        if len(hexpart) == 3:
            hexpart = "".join(c * 2 for c in hexpart)
        if len(hexpart) == 6 and re.fullmatch(r"[0-9a-f]{6}", hexpart):
            return hexpart.upper()
        return None
    if value in _NAMED_COLORS:
        return _NAMED_COLORS[value].upper()
    rgb_match = re.match(r"rgb\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)", value)
    if rgb_match:
        r, g, b = (int(x) for x in rgb_match.groups())
        return f"{r:02X}{g:02X}{b:02X}"
    return None


def _normalize_qt_html(html: str) -> str:
    """Rewrite Qt's span/margin-based "styles" into semantic HTML tags."""

    def heading_sub(match: "re.Match") -> str:
        top, bottom, style, inner = match.groups()
        if top != "12" or bottom != "12":
            return match.group(0)
        styles = _style_dict(style)
        size = styles.get("font-size", "")
        weight = styles.get("font-weight", "")
        is_bold = weight in ("700", "bold", "600")
        size_match = re.match(r"(\d+)pt", size)
        if not (is_bold and size_match):
            return match.group(0)
        level = _heading_level_for_size(int(size_match.group(1)))
        if level is None:
            return match.group(0)
        return f"<h{level}>{inner}</h{level}>"

    html = _HEADING_BLOCK_RE.sub(heading_sub, html)

    def span_sub(match: "re.Match") -> str:
        style, inner = match.groups()
        styles = _style_dict(style)
        weight = styles.get("font-weight", "")
        text_decoration = styles.get("text-decoration", "")
        color = styles.get("color", "").strip()
        background = styles.get("background-color", "").strip()

        # Font-family needs its original casing preserved (e.g. "Georgia"),
        # so pull it straight from the raw style string rather than the
        # lowercased _style_dict value.
        family_match = re.search(r"font-family:\s*([^;]+)", style, re.IGNORECASE)
        font_family_raw = family_match.group(1).strip() if family_match else ""
        font_family = font_family_raw.strip("'\"").lower()

        out = inner
        is_mono = any(mono in font_family for mono in _MONOSPACE_FAMILIES)
        if is_mono:
            out = f"<code>{out}</code>"
        elif font_family and _is_basic_font(font_family):
            out = f'<mdfont data="{font_family_raw.strip(chr(39)).strip(chr(34))}">{out}</mdfont>'

        if weight in ("700", "bold", "600"):
            out = f"<strong>{out}</strong>"
        if styles.get("font-style") == "italic":
            out = f"<em>{out}</em>"

        # Qt renders hyperlinks as blue + underlined; don't double that up
        # with a redundant color span on the link text itself.
        is_link_style = color == _LINK_BLUE and "underline" in text_decoration

        if "line-through" in text_decoration:
            out = f"<s>{out}</s>"
        elif "underline" in text_decoration and not is_link_style:
            out = f"<u>{out}</u>"

        if background and background not in ("transparent", "none"):
            out = f'<mdhl data="{background}">{out}</mdhl>'
        if color and not is_link_style and color not in ("#000000", "black"):
            out = f'<mdcolor data="{color}">{out}</mdcolor>'
        return out

    # Spans can't meaningfully nest in Qt's output, but run twice in case a
    # link's inner <span> left another plain span behind.
    for _ in range(2):
        new_html = _STYLE_ATTR_RE.sub(span_sub, html)
        if new_html == html:
            break
        html = new_html

    return html


# ============================================================================
# HTML -> Markdown
# ============================================================================

def strip_tags(text: str) -> str:
    """Remove any remaining HTML tags and unescape entities."""
    text = re.sub(r"<[^>]+>", "", text)
    return unescape(text).strip()


def _strip_tags_no_trim(text: str) -> str:
    """Like strip_tags, but keeps leading/trailing whitespace.

    Used for inline formatting (bold/italic/strike/code) so that, e.g.,
    "tender <strong>young</strong>" doesn't lose the space that separates
    the two runs.
    """
    text = re.sub(r"<[^>]+>", "", text)
    return unescape(text)


def _wrap_inline(inner_html: str, marker: str) -> str:
    """Wrap inner HTML's text in a Markdown emphasis marker.

    Markdown emphasis can't have whitespace touching the marker (CommonMark),
    so any leading/trailing space in *inner_html* is moved outside the
    markers rather than lost - "tender " + "young" stays "tender " when
    wrapped, instead of becoming "tender" and losing the word boundary.
    """
    text = _strip_tags_no_trim(inner_html)
    core = text.strip()
    if not core:
        return text
    leading = text[: len(text) - len(text.lstrip())]
    trailing = text[len(text.rstrip()):]
    return f"{leading}{marker}{core}{marker}{trailing}"


def _md_list_sub(list_tag: str, ordered: bool):
    item_re = re.compile(r"<li[^>]*>(.*?)</li>", re.IGNORECASE | re.DOTALL)

    def block_sub(match: "re.Match") -> str:
        inner = match.group(1)
        lines = []
        for index, item_match in enumerate(item_re.finditer(inner), start=1):
            text = strip_tags(item_match.group(1))
            if not text:
                continue
            prefix = f"{index}. " if ordered else "- "
            lines.append(prefix + text)
        return ("\n".join(lines) + "\n\n") if lines else ""

    return re.compile(rf"<{list_tag}[^>]*>(.*?)</{list_tag}>", re.IGNORECASE | re.DOTALL), block_sub


def _md_table_sub(match: "re.Match") -> str:
    table_html = match.group(0)
    row_re = re.compile(r"<tr[^>]*>(.*?)</tr>", re.IGNORECASE | re.DOTALL)
    cell_re = re.compile(r"<t[dh][^>]*>(.*?)</t[dh]>", re.IGNORECASE | re.DOTALL)

    rows = []
    for row_match in row_re.finditer(table_html):
        cells = [strip_tags(c).replace("|", "\\|") for c in cell_re.findall(row_match.group(1))]
        if cells:
            rows.append(cells)

    if not rows:
        return ""

    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]

    lines = ["| " + " | ".join(rows[0]) + " |", "| " + " | ".join(["---"] * width) + " |"]
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines) + "\n\n"


def html_to_markdown(html: str, html_fallback: bool = False) -> str:
    """Convert an HTML fragment/document into Markdown text.

    By default (``html_fallback=False``) the output is pure Markdown: things
    Markdown has no syntax for - underline, text color, highlighting, and
    non-basic fonts - are dropped, keeping only the plain text.

    With ``html_fallback=True``, those are instead emitted as inline HTML
    (``<u>``, ``<span style="color: ...">``, ``<span style="background-color:
    ...">``, ``<font face="...">``), which most Markdown renderers pass
    through unchanged.
    """
    text = _normalize_qt_html(html)

    # Drop everything outside the body, plus Qt's <style>/<head> boilerplate
    body_match = re.search(r"<body[^>]*>(.*)</body>", text, re.IGNORECASE | re.DOTALL)
    if body_match:
        text = body_match.group(1)
    text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r"<head[^>]*>.*?</head>", "", text, flags=re.IGNORECASE | re.DOTALL)

    # Tables first, so their <tr>/<td> guts aren't mangled by later rules
    text = re.sub(r"<table[^>]*>.*?</table>", _md_table_sub, text, flags=re.IGNORECASE | re.DOTALL)

    for tag, ordered in (("ul", False), ("ol", True)):
        pattern, sub_fn = _md_list_sub(tag, ordered)
        text = pattern.sub(sub_fn, text)

    text = re.sub(
        r"<h([1-6])[^>]*>(.*?)</h\1>",
        lambda m: "#" * int(m.group(1)) + " " + strip_tags(m.group(2)) + "\n\n",
        text, flags=re.IGNORECASE | re.DOTALL,
    )

    text = re.sub(r"<hr[^>]*/?>", "\n---\n\n", text, flags=re.IGNORECASE)

    def img_sub(match: "re.Match") -> str:
        attrs = match.group(1)
        src_match = re.search(r'src="([^"]*)"', attrs, re.IGNORECASE)
        alt_match = re.search(r'alt="([^"]*)"', attrs, re.IGNORECASE)
        src = src_match.group(1) if src_match else ""
        alt = alt_match.group(1) if alt_match else ""
        return f"![{alt}]({src})"

    text = re.sub(r"<img\s+([^>]*)/?>", img_sub, text, flags=re.IGNORECASE)

    text = re.sub(
        r'<a[^>]*\bhref="([^"]*)"[^>]*>(.*?)</a>',
        lambda m: f"[{strip_tags(m.group(2))}]({m.group(1)})",
        text, flags=re.IGNORECASE | re.DOTALL,
    )

    text = re.sub(
        r"<pre[^>]*>(.*?)</pre>",
        lambda m: "```\n" + strip_tags(m.group(1)) + "\n```\n\n",
        text, flags=re.IGNORECASE | re.DOTALL,
    )
    text = re.sub(
        r"<code[^>]*>(.*?)</code>",
        lambda m: _wrap_inline(m.group(1), "`"),
        text, flags=re.IGNORECASE | re.DOTALL,
    )

    text = re.sub(
        r"<(strong|b)[^>]*>(.*?)</\1>",
        lambda m: _wrap_inline(m.group(2), "**"),
        text, flags=re.IGNORECASE | re.DOTALL,
    )
    text = re.sub(
        r"<(em|i)[^>]*>(.*?)</\1>",
        lambda m: _wrap_inline(m.group(2), "*"),
        text, flags=re.IGNORECASE | re.DOTALL,
    )
    text = re.sub(
        r"<(s|strike|del)[^>]*>(.*?)</\1>",
        lambda m: _wrap_inline(m.group(2), "~~"),
        text, flags=re.IGNORECASE | re.DOTALL,
    )

    # Underline / color / highlight / font have no native Markdown syntax.
    # In plain mode we just drop the wrapper tags (the final strip_tags()
    # pass below takes care of that). In html_fallback mode we protect them
    # with \x00-delimited placeholders - which strip_tags() won't touch,
    # since it only matches literal '<...>' - and turn them back into real
    # HTML tags after strip_tags() has cleaned up everything else.
    if html_fallback:
        text = re.sub(
            r"<u>(.*?)</u>",
            lambda m: f"\x00U\x00{m.group(1)}\x00/U\x00",
            text, flags=re.IGNORECASE | re.DOTALL,
        )
        text = re.sub(
            r'<mdcolor data="([^"]*)">(.*?)</mdcolor>',
            lambda m: f"\x00C:{m.group(1)}\x00{m.group(2)}\x00/C\x00",
            text, flags=re.IGNORECASE | re.DOTALL,
        )
        text = re.sub(
            r'<mdhl data="([^"]*)">(.*?)</mdhl>',
            lambda m: f"\x00H:{m.group(1)}\x00{m.group(2)}\x00/H\x00",
            text, flags=re.IGNORECASE | re.DOTALL,
        )
        text = re.sub(
            r'<mdfont data="([^"]*)">(.*?)</mdfont>',
            lambda m: f"\x00F:{m.group(1)}\x00{m.group(2)}\x00/F\x00",
            text, flags=re.IGNORECASE | re.DOTALL,
        )

    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)

    text = re.sub(
        r"<p[^>]*>(.*?)</p>",
        lambda m: m.group(1).strip() + "\n\n",
        text, flags=re.IGNORECASE | re.DOTALL,
    )

    text = strip_tags(text)

    # Collapse runs of 3+ blank lines down to a single blank line
    text = re.sub(r"\n{3,}", "\n\n", text)

    if html_fallback:
        text = text.replace("\x00U\x00", "<u>").replace("\x00/U\x00", "</u>")
        text = re.sub(r"\x00C:([^\x00]*)\x00", lambda m: f'<span style="color: {m.group(1)};">', text)
        text = text.replace("\x00/C\x00", "</span>")
        text = re.sub(r"\x00H:([^\x00]*)\x00", lambda m: f'<span style="background-color: {m.group(1)};">', text)
        text = text.replace("\x00/H\x00", "</span>")
        text = re.sub(r"\x00F:([^\x00]*)\x00", lambda m: f'<font face="{m.group(1)}">', text)
        text = text.replace("\x00/F\x00", "</font>")

    return text.strip() + "\n"


# ============================================================================
# HTML -> DOCX
# ============================================================================

def _set_run_highlight(run, color_value: str):
    """Shade a run's background with an arbitrary RGB color (real Word highlight)."""
    hex_color = _to_hex(color_value)
    if not hex_color:
        return
    from docx.oxml.shared import OxmlElement, qn

    rpr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    rpr.append(shd)


def _set_run_font(run, font_name: str):
    """Set a run's font, including the east-asian fallback Word also checks."""
    from docx.oxml.shared import OxmlElement, qn

    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        rpr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def _extract_body_font(html: str) -> Optional[str]:
    """Get the document-wide default font Qt stores on <body style="...">.

    When a font is applied to the *whole* document (e.g. select-all then
    change font), Qt doesn't wrap every word in a <span> - it just sets the
    font once on <body style="font-family:...">, and individual runs simply
    inherit it. Without this, that common case would silently lose the font
    entirely, since the rest of the pipeline only looks at <span> styles.
    """
    body_match = re.search(r"<body[^>]*style=\"([^\"]*)\"", html, re.IGNORECASE)
    if not body_match:
        return None
    family_match = re.search(r"font-family:\s*([^;]+)", body_match.group(1), re.IGNORECASE)
    if not family_match:
        return None
    family = family_match.group(1).strip().strip("'\"")
    return family if _is_basic_font(family) else None


def _add_hyperlink(paragraph, url: str, text: str):
    """Insert a real, clickable hyperlink run into a python-docx paragraph."""
    from docx.oxml.shared import OxmlElement, qn

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_element = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_props.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(underline)

    run_element.append(run_props)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


class _HtmlToDocxParser(HTMLParser):
    """Streams normalized HTML into a python-docx ``Document``."""

    _HEADING_TAGS = {f"h{i}" for i in range(1, 7)}
    _INLINE_STYLE_TAGS = {
        "strong": "bold", "b": "bold",
        "em": "italic", "i": "italic",
        "u": "underline",
        "s": "strike", "strike": "strike", "del": "strike",
        "code": "code",
    }
    _WRAPPER_TAGS = {"mdcolor": "color", "mdhl": "highlight", "mdfont": "font"}

    def __init__(self, document, default_font: Optional[str] = None):
        super().__init__(convert_charrefs=True)
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        self._align_map = {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
        }
        self.document = document
        self.default_font = default_font
        self.current_paragraph = None
        # Each entry is a (kind, value) tuple, e.g. ("bold", None) or
        # ("color", "#ff0000"); value-bearing kinds use the most recently
        # pushed entry so nested tags of the same kind override correctly.
        self.style_stack: list[tuple[str, Optional[str]]] = []
        self.list_type: Optional[str] = None
        self.in_link = False
        self.link_href = ""
        self._link_text_parts: list[str] = []
        # Minimal table support: buffer rows/cells as plain text, then build
        # a real docx table once </table> is reached.
        self._in_table = False
        self._table_rows: list[list[str]] = []
        self._table_current_row: Optional[list[str]] = None
        self._table_cell_buffer: Optional[list[str]] = None

    # -- tag handling --------------------------------------------------

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)

        if tag == "table":
            self._in_table = True
            self._table_rows = []
            return
        if self._in_table:
            if tag == "tr":
                self._table_current_row = []
            elif tag in ("td", "th"):
                self._table_cell_buffer = []
            elif tag == "br" and self._table_cell_buffer is not None:
                self._table_cell_buffer.append("\n")
            return

        if tag in self._HEADING_TAGS:

            level = int(tag[1])
            self.current_paragraph = self.document.add_heading(level=level)
        elif tag == "p":
            self._new_paragraph()
            self._apply_paragraph_attrs(attrs_dict)
        elif tag in self._INLINE_STYLE_TAGS:
            self.style_stack.append((self._INLINE_STYLE_TAGS[tag], None))
        elif tag in self._WRAPPER_TAGS:
            self.style_stack.append((self._WRAPPER_TAGS[tag], attrs_dict.get("data", "")))
        elif tag == "a":
            self.in_link = True
            self.link_href = attrs_dict.get("href", "")
            self._link_text_parts = []
        elif tag == "ul":
            self.list_type = "ul"
        elif tag == "ol":
            self.list_type = "ol"
        elif tag == "li":
            style = "List Number" if self.list_type == "ol" else "List Bullet"
            self.current_paragraph = self.document.add_paragraph(style=style)
        elif tag == "br":
            if self.current_paragraph is not None:
                self.current_paragraph.add_run().add_break()
        elif tag == "img":
            self._add_image(attrs_dict.get("src", ""))
        elif tag == "hr":
            self._new_paragraph()
            self.current_paragraph.add_run("―" * 20)

    def handle_endtag(self, tag):
        if tag == "table":
            self._in_table = False
            self._build_table()
            return
        if self._in_table:
            if tag in ("td", "th") and self._table_cell_buffer is not None:
                if self._table_current_row is not None:
                    self._table_current_row.append("".join(self._table_cell_buffer).strip())
                self._table_cell_buffer = None
            elif tag == "tr" and self._table_current_row is not None:
                self._table_rows.append(self._table_current_row)
                self._table_current_row = None
            return

        if tag in self._HEADING_TAGS:
            self.current_paragraph = None
        elif tag == "p":
            self.current_paragraph = None
        elif tag in self._INLINE_STYLE_TAGS:
            self._pop_style(self._INLINE_STYLE_TAGS[tag])
        elif tag in self._WRAPPER_TAGS:
            self._pop_style(self._WRAPPER_TAGS[tag])
        elif tag == "a":
            text = "".join(self._link_text_parts).strip()
            if text and self.link_href:
                if self.current_paragraph is None:
                    self._new_paragraph()
                _add_hyperlink(self.current_paragraph, self.link_href, text)
            elif text:
                self._add_run(text)
            self.in_link = False
            self.link_href = ""
            self._link_text_parts = []
        elif tag in ("ul", "ol"):
            self.list_type = None
        elif tag == "li":
            self.current_paragraph = None

    def handle_data(self, data):
        if not data:
            return
        if self._in_table:
            if self._table_cell_buffer is not None:
                self._table_cell_buffer.append(data)
            return
        if self.in_link:
            self._link_text_parts.append(data)
            return
        if self.current_paragraph is None:
            if not data.strip():
                return
            self._new_paragraph()
        self._add_run(data)

    # -- helpers ---------------------------------------------------------

    def _pop_style(self, kind: str):
        """Remove the most recently pushed entry of *kind* (LIFO by tag)."""
        for i in range(len(self.style_stack) - 1, -1, -1):
            if self.style_stack[i][0] == kind:
                del self.style_stack[i]
                return

    def _new_paragraph(self):
        self.current_paragraph = self.document.add_paragraph()

    def _apply_paragraph_attrs(self, attrs_dict: dict):
        align = attrs_dict.get("align", "").strip().lower()
        if align in self._align_map:
            self.current_paragraph.alignment = self._align_map[align]

        style = attrs_dict.get("style", "")
        indent_match = re.search(r"-qt-block-indent:\s*(\d+)", style)
        if indent_match:
            level = int(indent_match.group(1))
            if level > 0:
                from docx.shared import Inches
                self.current_paragraph.paragraph_format.left_indent = Inches(0.5 * level)

        first_line_match = re.search(r"text-indent:\s*(-?\d+)px", style)
        if first_line_match:
            px = int(first_line_match.group(1))
            if px:
                from docx.shared import Pt
                self.current_paragraph.paragraph_format.first_line_indent = Pt(px * 0.75)

    def _add_run(self, text: str):
        if self.current_paragraph is None:
            self._new_paragraph()
        run = self.current_paragraph.add_run(text)

        kinds = {kind for kind, _ in self.style_stack}
        if "bold" in kinds:
            run.bold = True
        if "italic" in kinds:
            run.italic = True
        if "underline" in kinds:
            run.underline = True
        if "strike" in kinds:
            run.font.strike = True
        if "code" in kinds:
            _set_run_font(run, "Courier New")

        # Value-bearing styles: use the most recently pushed value of each kind.
        color_value = next((v for k, v in reversed(self.style_stack) if k == "color" and v), None)
        if color_value:
            hex_color = _to_hex(color_value)
            if hex_color:
                from docx.shared import RGBColor
                run.font.color.rgb = RGBColor.from_string(hex_color)

        font_value = next((v for k, v in reversed(self.style_stack) if k == "font" and v), None)
        if "code" in kinds:
            pass  # already handled above via _set_run_font(run, "Courier New")
        elif font_value:
            _set_run_font(run, font_value)
        elif self.default_font:
            _set_run_font(run, self.default_font)

        highlight_value = next((v for k, v in reversed(self.style_stack) if k == "highlight" and v), None)
        if highlight_value:
            _set_run_highlight(run, highlight_value)

        return run

    def _build_table(self):
        rows = [r for r in self._table_rows if r]
        if not rows:
            return
        cols = max(len(r) for r in rows)
        rows = [r + [""] * (cols - len(r)) for r in rows]

        table = self.document.add_table(rows=len(rows), cols=cols)
        try:
            table.style = "Light Grid Accent 1"
        except KeyError:
            pass
        for row_index, row in enumerate(rows):
            for col_index, cell_text in enumerate(row):
                cell = table.cell(row_index, col_index)
                cell.text = cell_text
                if self.default_font:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            _set_run_font(run, self.default_font)
        self.current_paragraph = None

    def _add_image(self, src: str):
        from docx.shared import Inches

        if not src.startswith("data:image/"):
            return
        try:
            header, b64data = src.split(",", 1)
            image_bytes = base64.b64decode(b64data)
        except (ValueError, base64.binascii.Error):
            return

        try:
            self._new_paragraph()
            run = self.current_paragraph.add_run()
            run.add_picture(BytesIO(image_bytes), width=Inches(5.5))
        except Exception:
            self._add_run("[image could not be embedded]")


def save_html_as_docx(html: str, filepath: Path):
    """Convert HTML into a .docx file and save it to *filepath*."""
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "The 'python-docx' package is required for Word export. "
            "Install it with: pip install python-docx"
        ) from exc

    filepath = Path(filepath)
    default_font = _extract_body_font(html)
    normalized = _normalize_qt_html(html)

    body_match = re.search(r"<body[^>]*>(.*)</body>", normalized, re.IGNORECASE | re.DOTALL)
    if body_match:
        normalized = body_match.group(1)
    normalized = re.sub(r"<style[^>]*>.*?</style>", "", normalized, flags=re.IGNORECASE | re.DOTALL)

    document = Document()

    # Word's default template adds ~10pt of space after every paragraph on
    # top of its actual content. Combined with a real blank paragraph for
    # each blank line in the source, that reads as two blank lines instead
    # of one. Zero it out so blank lines map 1:1 with what the user typed.
    from docx.shared import Pt
    normal_format = document.styles["Normal"].paragraph_format
    normal_format.space_before = Pt(0)
    normal_format.space_after = Pt(0)

    parser = _HtmlToDocxParser(document, default_font=default_font)
    parser.feed(normalized)
    parser.close()

    if not document.paragraphs and not document.tables:
        document.add_paragraph("")

    filepath.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(filepath))
