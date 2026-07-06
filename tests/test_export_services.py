# Tests for services/export_services.py: HTML -> Markdown and HTML -> DOCX
# conversion, including the Qt-flavoured markup produced by
# DocumentTab.get_content_html() (span-based headings/code, not real
# <h1>/<code> tags).
import base64

import pytest

from services.export_services import html_to_markdown, strip_tags, save_html_as_docx

docx = pytest.importorskip("docx")


# ============================================================================
# strip_tags
# ============================================================================

def test_strip_tags_removes_tags_and_unescapes_entities():
    assert strip_tags("<b>Hello &amp; welcome</b>") == "Hello & welcome"


# ============================================================================
# html_to_markdown - semantic HTML input
# ============================================================================

def test_html_to_markdown_basic_sections():
    html = "<h1>Title</h1><p>This is <strong>bold</strong> and <em>italic</em>.</p>"
    markdown = html_to_markdown(html)
    assert "# Title" in markdown
    assert "**bold**" in markdown
    assert "*italic*" in markdown


def test_html_to_markdown_links_and_images():
    html = '<p>Link: <a href="https://example.com">Example</a></p>'
    md = html_to_markdown(html)
    assert "[Example](https://example.com)" in md


def test_html_to_markdown_headings_all_levels():
    html = "".join(f"<h{i}>Heading {i}</h{i}>" for i in range(1, 7))
    md = html_to_markdown(html)
    for i in range(1, 7):
        assert f"{'#' * i} Heading {i}" in md


def test_html_to_markdown_lists():
    html = "<ul><li>First</li><li>Second</li></ul><ol><li>One</li><li>Two</li></ol>"
    md = html_to_markdown(html)
    assert "- First" in md
    assert "- Second" in md
    assert "1. One" in md
    assert "2. Two" in md


def test_html_to_markdown_table():
    html = "<table><tr><td>Name</td><td>Age</td></tr><tr><td>Alice</td><td>30</td></tr></table>"
    md = html_to_markdown(html)
    assert "| Name | Age |" in md
    assert "| Alice | 30 |" in md
    assert "| --- | --- |" in md


def test_html_to_markdown_code():
    html = "<p>Run <code>x = 1</code> please.</p><pre>line one\nline two</pre>"
    md = html_to_markdown(html)
    assert "`x = 1`" in md
    assert "```\nline one\nline two\n```" in md


def test_html_to_markdown_image():
    html = '<p><img src="picture.png" alt="A picture" /></p>'
    md = html_to_markdown(html)
    assert "![A picture](picture.png)" in md


def test_html_to_markdown_strikethrough():
    html = "<p>This is <s>wrong</s> right.</p>"
    md = html_to_markdown(html)
    assert "~~wrong~~" in md


# ============================================================================
# html_to_markdown - Qt-flavoured markup (span-based headings/code)
# ============================================================================

def test_html_to_markdown_qt_style_heading():
    html = (
        '<p style=" margin-top:12px; margin-bottom:12px;">'
        '<span style=" font-size:24pt; font-weight:700;">My Title</span></p>'
        '<p style=" margin-top:0px; margin-bottom:0px;">Body text.</p>'
    )
    md = html_to_markdown(html)
    assert "# My Title" in md
    assert "Body text." in md


def test_html_to_markdown_qt_style_heading_levels():
    sizes_to_level = {24: 1, 20: 2, 18: 3, 16: 4, 14: 5, 12: 6}
    for size, level in sizes_to_level.items():
        html = (
            '<p style=" margin-top:12px; margin-bottom:12px;">'
            f'<span style=" font-size:{size}pt; font-weight:700;">H{level}</span></p>'
        )
        md = html_to_markdown(html)
        assert f"{'#' * level} H{level}" in md


def test_html_to_markdown_qt_style_bold_italic_underline():
    html = (
        '<p>Normal <span style=" font-weight:700;">bold</span> and '
        '<span style=" font-style:italic;">italic</span> and '
        '<span style=" text-decoration: underline;">underline</span>.</p>'
    )
    md = html_to_markdown(html)
    assert "**bold**" in md
    assert "*italic*" in md
    assert "underline" in md  # no native markdown syntax, but text is kept


def test_html_to_markdown_qt_style_code_block():
    html = '<p><span style=" font-family:\'Courier New\';">print(1)</span></p>'
    md = html_to_markdown(html)
    assert "`print(1)`" in md


def test_html_to_markdown_qt_style_link():
    html = (
        '<p><a href="https://anthropic.com">'
        '<span style=" text-decoration: underline; color:#0000ff;">link</span></a></p>'
    )
    md = html_to_markdown(html)
    assert "[link](https://anthropic.com)" in md


def test_html_to_markdown_does_not_leave_stray_tags():
    html = "<h1>Title</h1><p>Some <strong>bold</strong> text with a <br/>break.</p>"
    md = html_to_markdown(html)
    assert "<" not in md and ">" not in md


# ============================================================================
# save_html_as_docx
# ============================================================================

def test_save_html_as_docx_basic(tmp_path):
    from docx import Document

    html = "<h1>Report</h1><p>Some <strong>bold</strong> and <em>italic</em> text.</p>"
    out_path = tmp_path / "out.docx"
    save_html_as_docx(html, out_path)

    assert out_path.exists()
    doc = Document(str(out_path))
    texts = [p.text for p in doc.paragraphs]
    assert "Report" in texts
    assert any("bold" in t and "italic" in t for t in texts)


def test_save_html_as_docx_heading_style(tmp_path):
    from docx import Document

    html = "<h1>Top</h1><h2>Sub</h2>"
    out_path = tmp_path / "headings.docx"
    save_html_as_docx(html, out_path)

    doc = Document(str(out_path))
    styles = {p.text: p.style.name for p in doc.paragraphs}
    assert styles.get("Top") == "Heading 1"
    assert styles.get("Sub") == "Heading 2"


def test_save_html_as_docx_lists(tmp_path):
    from docx import Document

    html = "<ul><li>Alpha</li><li>Beta</li></ul>"
    out_path = tmp_path / "list.docx"
    save_html_as_docx(html, out_path)

    doc = Document(str(out_path))
    bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert [p.text for p in bullets] == ["Alpha", "Beta"]


def test_save_html_as_docx_table(tmp_path):
    from docx import Document

    html = "<table><tr><td>A</td><td>B</td></tr><tr><td>1</td><td>2</td></tr></table>"
    out_path = tmp_path / "table.docx"
    save_html_as_docx(html, out_path)

    doc = Document(str(out_path))
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert [c.text for c in table.rows[0].cells] == ["A", "B"]
    assert [c.text for c in table.rows[1].cells] == ["1", "2"]


def test_save_html_as_docx_hyperlink(tmp_path):
    from docx import Document

    html = '<p>Visit <a href="https://example.com">Example</a> now.</p>'
    out_path = tmp_path / "link.docx"
    save_html_as_docx(html, out_path)

    doc = Document(str(out_path))
    # A real w:hyperlink element should be present in the paragraph XML.
    xml = doc.paragraphs[0]._p.xml
    assert "w:hyperlink" in xml
    assert "Example" in xml


def test_save_html_as_docx_image(tmp_path):
    from docx import Document

    png_1x1 = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk"
        "+A8AAQUBAScY42YAAAAASUVORK5CYII="
    )
    b64 = base64.b64encode(png_1x1).decode()
    html = f'<p><img src="data:image/png;base64,{b64}" /></p>'
    out_path = tmp_path / "image.docx"
    save_html_as_docx(html, out_path)

    doc = Document(str(out_path))
    assert len(doc.inline_shapes) == 1


def test_save_html_as_docx_creates_parent_dirs(tmp_path):
    out_path = tmp_path / "nested" / "dir" / "out.docx"
    save_html_as_docx("<p>Hello</p>", out_path)
    assert out_path.exists()


def test_save_html_as_docx_empty_document_does_not_crash(tmp_path):
    out_path = tmp_path / "empty.docx"
    save_html_as_docx('<p style="-qt-paragraph-type:empty;"><br /></p>', out_path)
    assert out_path.exists()


# ============================================================================
# html_to_markdown - html_fallback mode
# ============================================================================

def test_html_to_markdown_plain_drops_unsupported_formatting():
    html = (
        '<p>Normal <span style=" text-decoration: underline;">underline</span> and '
        '<span style=" color:#ff0000;">red</span> and '
        '<span style=" background-color:#ffff00;">highlighted</span>.</p>'
    )
    md = html_to_markdown(html)
    assert "<" not in md and ">" not in md
    assert "underline" in md and "red" in md and "highlighted" in md


def test_html_to_markdown_fallback_keeps_underline_as_html():
    html = '<p><span style=" text-decoration: underline;">underline</span></p>'
    md = html_to_markdown(html, html_fallback=True)
    assert "<u>underline</u>" in md


def test_html_to_markdown_fallback_keeps_color_as_span():
    html = '<p><span style=" color:#ff0000;">red text</span></p>'
    md = html_to_markdown(html, html_fallback=True)
    assert '<span style="color: #ff0000;">red text</span>' in md


def test_html_to_markdown_fallback_keeps_highlight_as_span():
    html = '<p><span style=" background-color:#ffff00;">highlighted</span></p>'
    md = html_to_markdown(html, html_fallback=True)
    assert '<span style="background-color: #ffff00;">highlighted</span>' in md


def test_html_to_markdown_fallback_basic_font_uses_font_tag():
    html = '<p><span style=" font-family:\'Georgia\';">georgia text</span></p>'
    md = html_to_markdown(html, html_fallback=True)
    assert '<font face="Georgia">georgia text</font>' in md


def test_html_to_markdown_non_basic_font_falls_back_to_plain_text():
    html = '<p><span style=" font-family:\'SomeExoticFont\';">exotic text</span></p>'
    for fallback in (False, True):
        md = html_to_markdown(html, html_fallback=fallback)
        assert "exotic text" in md
        assert "<font" not in md


def test_html_to_markdown_preserves_spacing_between_adjacent_inline_runs():
    html = "<p><em>tender </em><strong>young thing!</strong></p>"
    md = html_to_markdown(html)
    assert "*tender* **young thing!**" in md


def test_html_to_markdown_link_not_recolored_by_qt_link_style():
    html = (
        '<p><a href="https://example.com">'
        '<span style=" text-decoration: underline; color:#0000ff;">link</span></a></p>'
    )
    md = html_to_markdown(html, html_fallback=True)
    assert "[link](https://example.com)" in md
    assert "<span" not in md


# ============================================================================
# save_html_as_docx - color, highlight, font, alignment, indent
# ============================================================================

def test_save_html_as_docx_text_color(tmp_path):
    from docx import Document

    html = '<p><span style=" color:#ff0000;">Red text</span></p>'
    out_path = tmp_path / "color.docx"
    save_html_as_docx(html, out_path)

    doc = Document(str(out_path))
    run = doc.paragraphs[0].runs[0]
    assert str(run.font.color.rgb) == "FF0000"


def test_save_html_as_docx_highlight(tmp_path):
    from docx import Document

    html = '<p><span style=" background-color:#ffff00;">Highlighted</span></p>'
    out_path = tmp_path / "highlight.docx"
    save_html_as_docx(html, out_path)

    doc = Document(str(out_path))
    run_xml = doc.paragraphs[0].runs[0]._element.xml
    assert 'w:fill="FFFF00"' in run_xml


def test_save_html_as_docx_basic_font_applied(tmp_path):
    from docx import Document

    html = '<p><span style=" font-family:\'Georgia\';">Georgia text</span></p>'
    out_path = tmp_path / "font.docx"
    save_html_as_docx(html, out_path)

    doc = Document(str(out_path))
    assert doc.paragraphs[0].runs[0].font.name == "Georgia"


def test_save_html_as_docx_non_basic_font_falls_back_to_normal(tmp_path):
    from docx import Document

    html = '<p><span style=" font-family:\'SomeExoticFont\';">Exotic text</span></p>'
    out_path = tmp_path / "exotic_font.docx"
    save_html_as_docx(html, out_path)

    doc = Document(str(out_path))
    assert doc.paragraphs[0].runs[0].font.name is None


def test_save_html_as_docx_alignment(tmp_path):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    html = '<p align="center">Centered</p><p align="right">Right</p>'
    out_path = tmp_path / "alignment.docx"
    save_html_as_docx(html, out_path)

    doc = Document(str(out_path))
    assert doc.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert doc.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.RIGHT


def test_save_html_as_docx_indent(tmp_path):
    from docx import Document
    from docx.shared import Inches

    html = '<p style=" -qt-block-indent:2; text-indent:0px;">Indented</p>'
    out_path = tmp_path / "indent.docx"
    save_html_as_docx(html, out_path)

    doc = Document(str(out_path))
    assert doc.paragraphs[0].paragraph_format.left_indent == Inches(1.0)


def test_save_html_as_docx_color_and_highlight_together(tmp_path):
    from docx import Document

    html = '<p><span style=" color:#ff0000; background-color:#ffff00;">Both</span></p>'
    out_path = tmp_path / "combo.docx"
    save_html_as_docx(html, out_path)

    doc = Document(str(out_path))
    run = doc.paragraphs[0].runs[0]
    assert str(run.font.color.rgb) == "FF0000"
    assert 'w:fill="FFFF00"' in run._element.xml


def test_save_html_as_docx_whole_document_font_from_body_style(tmp_path):
    """A font set on <body style="font-family:...;"> (Qt's representation
    of a document-wide font change, e.g. select-all then change font) should
    apply to ordinary runs that have no per-span font of their own."""
    from docx import Document

    html = (
        '<html><body style=" font-family:\'Georgia\'; font-size:12pt;">'
        "<p>Plain paragraph text.</p>"
        "</body></html>"
    )
    out_path = tmp_path / "body_font.docx"
    save_html_as_docx(html, out_path)

    doc = Document(str(out_path))
    assert doc.paragraphs[0].runs[0].font.name == "Georgia"


def test_save_html_as_docx_body_font_overridden_by_explicit_span(tmp_path):
    from docx import Document

    html = (
        '<html><body style=" font-family:\'Georgia\';">'
        '<p>Georgia text <span style=" font-family:\'Arial\';">Arial text</span></p>'
        "</body></html>"
    )
    out_path = tmp_path / "body_font_override.docx"
    save_html_as_docx(html, out_path)

    doc = Document(str(out_path))
    runs = doc.paragraphs[0].runs
    assert runs[0].font.name == "Georgia"
    assert runs[-1].font.name == "Arial"


def test_save_html_as_docx_non_basic_body_font_falls_back_to_normal(tmp_path):
    from docx import Document

    html = (
        '<html><body style=" font-family:\'SomeExoticFont\';">'
        "<p>Plain text.</p></body></html>"
    )
    out_path = tmp_path / "exotic_body_font.docx"
    save_html_as_docx(html, out_path)

    doc = Document(str(out_path))
    assert doc.paragraphs[0].runs[0].font.name is None


def test_save_html_as_docx_zero_paragraph_spacing(tmp_path):
    """Word's default template adds ~10pt of space after every paragraph;
    with a real blank paragraph for each blank line in the source, that
    doubled up as two blank lines' worth of gap. Should be zeroed out."""
    from docx import Document
    from docx.shared import Pt

    out_path = tmp_path / "spacing.docx"
    save_html_as_docx("<p>Hello</p>", out_path)

    doc = Document(str(out_path))
    normal_format = doc.styles["Normal"].paragraph_format
    assert normal_format.space_after == Pt(0)
    assert normal_format.space_before == Pt(0)


def test_save_html_as_docx_single_blank_line_stays_single(tmp_path):
    """One blank line in the source should produce exactly one empty
    paragraph in the output, not extra ones."""
    from docx import Document

    html = (
        "<p>First paragraph.</p>"
        '<p style="-qt-paragraph-type:empty;"><br /></p>'
        "<p>Second paragraph.</p>"
    )
    out_path = tmp_path / "blank_line.docx"
    save_html_as_docx(html, out_path)

    doc = Document(str(out_path))
    texts = [p.text for p in doc.paragraphs]
    assert texts == ["First paragraph.", "\n", "Second paragraph."]
